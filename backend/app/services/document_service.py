import os
import hashlib
import uuid
from typing import List, Optional, BinaryIO
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
import aiofiles

from app.core.config import settings
from app.core.logging import get_logger
from app.models.document import Document, DocumentChunk
from app.services.llm_service import llm_service

logger = get_logger(__name__)


class DocumentService:
    """Service for document management and processing."""
    
    def __init__(self):
        self.storage_path = settings.DOCUMENT_STORAGE_PATH
        self.max_size_mb = settings.MAX_DOCUMENT_SIZE_MB
        self.supported_types = settings.SUPPORTED_DOCUMENT_TYPES
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    def _get_file_path(self, document_id: str, filename: str) -> str:
        """Get the storage path for a document."""
        # Create subdirectory based on document ID prefix for better organization
        subdir = document_id[:2]
        dir_path = os.path.join(self.storage_path, subdir)
        os.makedirs(dir_path, exist_ok=True)
        return os.path.join(dir_path, f"{document_id}_{filename}")
    
    def _calculate_checksum(self, content: bytes) -> str:
        """Calculate SHA-256 checksum of content."""
        return hashlib.sha256(content).hexdigest()
    
    async def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text content from a document."""
        
        if mime_type == "text/plain":
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return await f.read()
        
        elif mime_type == "text/markdown":
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return await f.read()
        
        elif mime_type == "application/pdf":
            try:
                import pypdf
                text = ""
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n\n"
                return text
            except Exception as e:
                logger.error(f"PDF extraction error: {e}")
                return ""
        
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                return ""
        
        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                text = ""
                for sheet in wb.worksheets:
                    text += f"\n\nSheet: {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) for cell in row if cell is not None])
                        if row_text.strip():
                            text += row_text + "\n"
                return text
            except Exception as e:
                logger.error(f"XLSX extraction error: {e}")
                return ""
        
        else:
            logger.warning(f"Unsupported MIME type for text extraction: {mime_type}")
            return ""
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at a sentence or paragraph boundary
            if end < len(text):
                # Look for sentence boundary
                for i in range(min(end, len(text) - 1), start + self.chunk_size // 2, -1):
                    if text[i] in '.!?\n':
                        end = i + 1
                        chunk = text[start:end]
                        break
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return chunks
    
    async def upload_document(
        self,
        db: AsyncSession,
        file_content: bytes,
        filename: str,
        mime_type: str,
        owner_id: str,
        title: Optional[str] = None,
        description: Optional[str] = None,
        tags: Optional[List[str]] = None,
        team_access: Optional[List[str]] = None,
    ) -> Document:
        """Upload and process a new document."""
        
        # Validate file size
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > self.max_size_mb:
            raise ValueError(f"File size ({file_size_mb:.1f}MB) exceeds maximum ({self.max_size_mb}MB)")
        
        # Validate MIME type
        if mime_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        # Calculate checksum
        checksum = self._calculate_checksum(file_content)
        
        # Check for duplicate
        existing = await db.execute(
            select(Document).where(Document.checksum == checksum)
        )
        if existing.scalar_one_or_none():
            raise ValueError("Document with identical content already exists")
        
        # Create document record
        document_id = str(uuid.uuid4())
        file_path = self._get_file_path(document_id, filename)
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        document = Document(
            id=document_id,
            title=title or filename,
            description=description,
            filename=filename,
            file_path=file_path,
            file_size=len(file_content),
            mime_type=mime_type,
            checksum=checksum,
            owner_id=owner_id,
            tags=tags or [],
            team_access=team_access or [],
            status="pending",
        )
        
        db.add(document)
        await db.commit()
        await db.refresh(document)
        
        logger.info(f"Document uploaded: {document_id} - {filename}")
        
        return document
    
    async def process_document(
        self,
        db: AsyncSession,
        document_id: str,
    ) -> Document:
        """Process a document: extract text, chunk, and generate embeddings."""
        
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            raise ValueError(f"Document not found: {document_id}")
        
        try:
            # Update status
            document.status = "processing"
            await db.commit()
            
            # Extract text
            content = await self._extract_text(document.file_path, document.mime_type)
            document.content = content
            
            if not content.strip():
                raise ValueError("No text content could be extracted from document")
            
            # Chunk the text
            chunks = self._chunk_text(content)
            
            # Generate embeddings for chunks
            chunk_objects = []
            for i, chunk_text in enumerate(chunks):
                # Get embedding
                embeddings = await llm_service.get_embeddings([chunk_text])
                embedding = embeddings[0] if embeddings else None
                
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=i,
                    embedding=embedding,
                    metadata_json={
                        "start_char": i * (self.chunk_size - self.chunk_overlap),
                        "end_char": i * (self.chunk_size - self.chunk_overlap) + len(chunk_text),
                    },
                )
                chunk_objects.append(chunk)
            
            # Save chunks
            for chunk in chunk_objects:
                db.add(chunk)
            
            # Update document
            document.status = "completed"
            document.chunk_count = len(chunks)
            document.processed_at = datetime.utcnow()
            
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Document processed: {document_id} - {len(chunks)} chunks")
            
            return document
        
        except Exception as e:
            document.status = "failed"
            document.processing_error = str(e)
            await db.commit()
            
            logger.error(f"Document processing failed: {document_id} - {e}")
            raise
    
    async def get_document(
        self,
        db: AsyncSession,
        document_id: str,
    ) -> Optional[Document]:
        """Get a document by ID."""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        return result.scalar_one_or_none()
    
    async def list_documents(
        self,
        db: AsyncSession,
        skip: int = 0,
        limit: int = 100,
        filters: Optional[dict] = None,
    ) -> List[Document]:
        """List documents with optional filtering."""
        query = select(Document)
        
        if filters:
            if "status" in filters:
                query = query.where(Document.status == filters["status"])
            if "owner_id" in filters:
                query = query.where(Document.owner_id == filters["owner_id"])
            if "tags" in filters:
                query = query.where(Document.tags.overlap(filters["tags"]))
        
        query = query.offset(skip).limit(limit).order_by(Document.created_at.desc())
        
        result = await db.execute(query)
        return result.scalars().all()
    
    async def delete_document(
        self,
        db: AsyncSession,
        document_id: str,
    ) -> bool:
        """Delete a document and its associated files."""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return False
        
        # Delete file
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as e:
            logger.error(f"Error deleting file: {e}")
        
        # Delete database record (cascades to chunks)
        await db.delete(document)
        await db.commit()
        
        logger.info(f"Document deleted: {document_id}")
        return True


# Global instance
document_service = DocumentService()
